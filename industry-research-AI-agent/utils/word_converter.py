# utils/word_converter.py
"""
Markdown转Word转换器
支持中文字体、表格渲染、专业排版
"""

import os
import re
from datetime import datetime
from typing import Optional
from io import BytesIO

# 尝试导入python-docx库
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def parse_markdown_table(table_text: str) -> list:
    """解析Markdown表格"""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return []
    
    rows = []
    for line in lines:
        if line.startswith('|') and line.endswith('|'):
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    return rows


def convert_md_to_word(
    md_content: str,
    output_path: Optional[str] = None,
    title: str = "行业研究报告",
    add_cover: bool = True,
    province: str = "",
    industry: str = "",
    year: str = ""
) -> bytes:
    """
    将Markdown内容转换为Word文档
    
    Args:
        md_content: Markdown文本内容
        output_path: 输出文件路径（可选）
        title: 报告标题
        add_cover: 是否添加封面
        province: 省份
        industry: 行业
        year: 年份
    
    Returns:
        Word文件的字节内容
    """
    if not HAS_DOCX:
        raise ImportError("需要安装python-docx库: pip install python-docx")
    
    # 创建文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
    
    # 定义样式
    styles = doc.styles
    
    # 标题1样式
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Microsoft YaHei'
    style_h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h1.font.size = Pt(22)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    
    # 标题2样式
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Microsoft YaHei'
    style_h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h2.font.size = Pt(16)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    
    # 标题3样式
    style_h3 = styles['Heading 3']
    style_h3.font.name = 'Microsoft YaHei'
    style_h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True
    style_h3.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
    
    # 正文样式
    style_normal = styles['Normal']
    style_normal.font.name = 'Microsoft YaHei'
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.first_line_indent = Cm(0.74)
    
    # 添加封面
    if add_cover:
        report_title = f"{year}年{province}{industry}行业研究报告" if province and industry else title
        
        # 封面标题
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.space_before = Pt(200)
        run = cover_title.add_run(report_title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("深度行业研究报告")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 报告信息
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.paragraph_format.space_before = Pt(100)
        run = info.add_run(f"FinSight AI Agent\n报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 免责声明
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer.paragraph_format.space_before = Pt(50)
        run = disclaimer.add_run("本报告由AI智能体自动生成，仅供参考")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 分页
        doc.add_page_break()
    
    # 解析Markdown内容
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                # 开始代码块
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 表格处理
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # 表格结束，渲染表格
            rows = parse_markdown_table('\n'.join(table_lines))
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text.strip()
                        
                        # 设置单元格样式
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Microsoft YaHei'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                run.font.size = Pt(10)
                        
                        # 表头样式
                        if row_idx == 0:
                            set_cell_shading(cell, '1e40af')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                        elif row_idx % 2 == 0:
                            set_cell_shading(cell, 'f8fafc')
                
                doc.add_paragraph()  # 表格后空行
            
            in_table = False
            table_lines = []
            # 不增加i，继续处理当前行
        
        # 标题处理
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(line[5:].strip())
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 列表处理
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.first_line_indent = Cm(0)
            
            # 处理粗体
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        elif line.strip() and re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(text)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 引用块
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 分割线
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run('─' * 50)
            run.font.color.rgb = RGBColor(0xe5, 0xe7, 0xeb)
        
        # 普通段落
        elif line.strip():
            p = doc.add_paragraph()
            
            # 处理粗体和斜体
            text = line.strip()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.font.italic = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        i += 1
    
    # 处理最后可能未结束的表格
    if in_table and table_lines:
        rows = parse_markdown_table('\n'.join(table_lines))
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_text.strip()
    
    # 保存到字节流
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    word_bytes = file_stream.read()
    
    # 如果指定了输出路径，保存文件
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(word_bytes)
    
    return word_bytes


def convert_md_file_to_word(
    md_file_path: str,
    output_path: Optional[str] = None,
    **kwargs
) -> bytes:
    """
    将Markdown文件转换为Word文档
    
    Args:
        md_file_path: Markdown文件路径
        output_path: 输出Word文件路径（可选，默认同名.docx）
        **kwargs: 传递给convert_md_to_word的其他参数
    
    Returns:
        Word文件的字节内容
    """
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 如果没有指定输出路径，使用同名.docx
    if not output_path:
        output_path = md_file_path.rsplit('.', 1)[0] + '.docx'
    
    # 从文件名提取信息
    filename = os.path.basename(md_file_path)
    parts = filename.replace('.md', '').split('_')
    
    # 尝试解析文件名中的年份、省份、行业
    year = kwargs.get('year', '')
    province = kwargs.get('province', '')
    industry = kwargs.get('industry', '')
    
    if len(parts) >= 3 and not year:
        year = parts[0] if parts[0].isdigit() else ''
    if len(parts) >= 3 and not province:
        province = parts[1] if '省' in parts[1] or '市' in parts[1] else ''
    if len(parts) >= 3 and not industry:
        industry = parts[2] if parts[2] else ''
    
    return convert_md_to_word(
        md_content=md_content,
        output_path=output_path,
        year=year,
        province=province,
        industry=industry,
        **kwargs
    )


if __name__ == "__main__":
    # 测试代码
    test_md = """
# 测试报告

## 第一章 概述

这是一个测试段落，用于验证Word转换功能。

### 1.1 背景

| 指标 | 数值 | 说明 |
|------|------|------|
| 市场规模 | 1000亿 | 2024年 |
| 增长率 | 15% | 年均 |

## 第二章 分析

**重要结论**：这是一个重要的发现。

- 要点1
- 要点2
- 要点3

> 这是一段引用文本
    """
    
    try:
        word_bytes = convert_md_to_word(
            test_md,
            output_path="test_report.docx",
            title="测试报告",
            province="浙江省",
            industry="人工智能",
            year="2025"
        )
        print(f"Word生成成功，大小: {len(word_bytes)} bytes")
    except Exception as e:
        print(f"Word生成失败: {e}")
